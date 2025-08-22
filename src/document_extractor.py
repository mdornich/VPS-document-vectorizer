import io
import json
import gc
import sys
from typing import Dict, Any, List, Optional, Union
from pathlib import Path
import tempfile
import mimetypes

import pandas as pd
import pypdf
from docx import Document
import mammoth
import openpyxl
import chardet
import structlog

from config.settings import settings

logger = structlog.get_logger()

# File size limits (in bytes)
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
MAX_PDF_PAGES = 1000  # Maximum pages to process in a PDF
MAX_EXCEL_ROWS = 100000  # Maximum rows to process in Excel files


class DocumentExtractor:
    """Extract text and structured data from various document formats."""
    
    # Supported MIME types and their handlers
    MIME_HANDLERS = {
        'application/pdf': 'extract_pdf',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'extract_docx',
        'application/msword': 'extract_doc',
        'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': 'extract_excel',
        'application/vnd.ms-excel': 'extract_excel',
        'text/csv': 'extract_csv',
        'text/plain': 'extract_text',
        'application/vnd.google-apps.document': 'extract_text',  # Already converted
        'application/vnd.google-apps.spreadsheet': 'extract_csv',  # Already converted
    }
    
    def __init__(self):
        self.temp_dir = Path(settings.temp_download_dir)
        self.temp_dir.mkdir(parents=True, exist_ok=True)
    
    def extract(
        self, 
        content: bytes, 
        file_metadata: Dict[str, Any]
    ) -> Dict[str, Any]:
        """
        Extract content from a document.
        
        Args:
            content: File content as bytes
            file_metadata: Metadata about the file
        
        Returns:
            Dictionary with extracted content and metadata
        """
        mime_type = file_metadata.get('mimeType', '')
        file_name = file_metadata.get('name', 'unknown')
        file_size = len(content)
        
        logger.info(f"Extracting content from {file_name} (type: {mime_type}, size: {file_size:,} bytes)")
        
        # Check file size limits
        if file_size > MAX_FILE_SIZE:
            error_msg = f"File too large: {file_size:,} bytes (max: {MAX_FILE_SIZE:,} bytes)"
            logger.warning(error_msg)
            return self._create_error_result(error_msg)
        
        # Log memory usage before processing
        self._log_memory_usage(f"Before processing {file_name}")
        
        # Find appropriate handler
        handler_name = self.MIME_HANDLERS.get(mime_type)
        
        if not handler_name:
            # Try to detect from file extension
            mime_type = self._detect_mime_type(file_name, content)
            handler_name = self.MIME_HANDLERS.get(mime_type)
        
        if not handler_name:
            logger.warning(f"Unsupported file type: {mime_type}")
            return self._create_error_result(f"Unsupported file type: {mime_type}")
        
        # Get handler method
        handler = getattr(self, handler_name, None)
        if not handler:
            logger.error(f"Handler {handler_name} not implemented")
            return self._create_error_result(f"Handler not implemented: {handler_name}")
        
        try:
            # Execute extraction
            result = handler(content, file_metadata)
            result['mime_type'] = mime_type
            result['file_name'] = file_name
            result['file_size'] = file_size
            
            # Force garbage collection after processing
            gc.collect()
            self._log_memory_usage(f"After processing {file_name}")
            
            return result
            
        except Exception as e:
            logger.error(f"Error extracting {file_name}: {e}")
            # Force cleanup on error
            gc.collect()
            return self._create_error_result(str(e))
    
    def extract_pdf(self, content: bytes, metadata: Dict[str, Any]) -> Dict[str, Any]:
        """Extract text from PDF files with page limits."""
        pdf_file = None
        try:
            pdf_file = io.BytesIO(content)
            pdf_reader = pypdf.PdfReader(pdf_file)
            
            total_pages = len(pdf_reader.pages)
            pages_to_process = min(total_pages, MAX_PDF_PAGES)
            
            if total_pages > MAX_PDF_PAGES:
                logger.warning(f"PDF has {total_pages} pages, processing only first {MAX_PDF_PAGES}")
            
            text_content = []
            for page_num in range(pages_to_process):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                if page_text:
                    text_content.append(page_text)
                
                # Periodic garbage collection for large PDFs
                if page_num % 50 == 0 and page_num > 0:
                    gc.collect()
            
            full_text = "\n\n".join(text_content)
            
            return {
                'type': 'text',
                'content': full_text,
                'page_count': total_pages,
                'pages_processed': pages_to_process,
                'extraction_method': 'pypdf_limited'
            }
            
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            raise
        finally:
            # Explicit cleanup
            if pdf_file:
                pdf_file.close()
            del content  # Help garbage collector
    
    def extract_docx(self, content: bytes, metadata: Dict[str, Any]) -> Dict[str, Any]:
        """Extract text from DOCX files using multiple methods."""
        try:
            # Method 1: Try mammoth first (better formatting preservation)
            try:
                result = mammoth.extract_raw_text(io.BytesIO(content))
                if result.value and len(result.value.strip()) > 0:
                    logger.info("Successfully extracted DOCX using mammoth")
                    return {
                        'type': 'text',
                        'content': result.value,
                        'extraction_method': 'mammoth',
                        'messages': result.messages
                    }
            except Exception as e:
                logger.debug(f"Mammoth extraction failed, trying python-docx: {e}")
            
            # Method 2: Fall back to python-docx
            doc = Document(io.BytesIO(content))
            
            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Extract tables
            tables_text = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    if any(row_data):
                        table_data.append(" | ".join(row_data))
                if table_data:
                    tables_text.append("\n".join(table_data))
            
            # Combine all text
            full_text = "\n\n".join(paragraphs)
            if tables_text:
                full_text += "\n\n" + "\n\n".join(tables_text)
            
            return {
                'type': 'text',
                'content': full_text,
                'paragraph_count': len(paragraphs),
                'table_count': len(doc.tables),
                'extraction_method': 'python-docx'
            }
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise
    
    def extract_doc(self, content: bytes, metadata: Dict[str, Any]) -> Dict[str, Any]:
        """Extract text from old DOC files."""
        try:
            # Try mammoth for DOC files too
            with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as tmp_file:
                tmp_file.write(content)
                tmp_path = tmp_file.name
            
            try:
                with open(tmp_path, 'rb') as doc_file:
                    result = mammoth.extract_raw_text(doc_file)
                    
                return {
                    'type': 'text',
                    'content': result.value,
                    'extraction_method': 'mammoth_doc',
                    'messages': result.messages
                }
            finally:
                Path(tmp_path).unlink(missing_ok=True)
                
        except Exception as e:
            logger.error(f"DOC extraction failed: {e}")
            # Try as text fallback
            return self.extract_text(content, metadata)
    
    def extract_excel(self, content: bytes, metadata: Dict[str, Any]) -> Dict[str, Any]:
        """Extract data from Excel files with row limits."""
        excel_file = None
        try:
            excel_file = io.BytesIO(content)
            
            # Read all sheets with better NaN handling and row limits
            all_sheets = pd.read_excel(
                excel_file, 
                sheet_name=None, 
                na_filter=True,
                nrows=MAX_EXCEL_ROWS  # Limit rows per sheet
            )
            
            structured_data = []
            text_content = []
            schema = {}
            total_rows = 0
            
            for sheet_name, df in all_sheets.items():
                # Check total row limit across all sheets
                if total_rows >= MAX_EXCEL_ROWS:
                    logger.warning(f"Excel row limit reached ({MAX_EXCEL_ROWS}), skipping remaining sheets")
                    break
                
                # Limit rows if needed
                remaining_rows = MAX_EXCEL_ROWS - total_rows
                if len(df) > remaining_rows:
                    df = df.head(remaining_rows)
                    logger.warning(f"Sheet '{sheet_name}' truncated to {remaining_rows} rows")
                
                # Replace NaN values with None for better JSON serialization
                df = df.where(pd.notnull(df), None)
                
                # Store schema
                schema[sheet_name] = df.columns.tolist()
                
                # Convert to records with None instead of NaN
                records = df.to_dict('records')
                structured_data.extend(records)
                total_rows += len(df)
                
                # Create text representation (limited to avoid memory issues)
                text_content.append(f"Sheet: {sheet_name} ({len(df)} rows)")
                if len(df) <= 1000:  # Only show full data for small sheets
                    text_content.append(df.to_string())
                else:
                    text_content.append(f"[Large sheet with {len(df)} rows - showing first 10 rows]")
                    text_content.append(df.head(10).to_string())
                
                # Periodic garbage collection
                if total_rows % 10000 == 0 and total_rows > 0:
                    gc.collect()
            
            return {
                'type': 'structured',
                'content': "\n\n".join(text_content),
                'data': structured_data,
                'schema': schema,
                'sheet_count': len(all_sheets),
                'row_count': total_rows,
                'extraction_method': 'pandas_excel_limited'
            }
            
        except Exception as e:
            logger.error(f"Excel extraction failed: {e}")
            raise
        finally:
            # Explicit cleanup
            if excel_file:
                excel_file.close()
            del content  # Help garbage collector
    
    def extract_csv(self, content: bytes, metadata: Dict[str, Any]) -> Dict[str, Any]:
        """Extract data from CSV files with row limits."""
        try:
            # Detect encoding
            detected = chardet.detect(content[:10000])
            encoding = detected.get('encoding', 'utf-8')
            
            # Read CSV with row limit
            csv_text = content.decode(encoding)
            df = pd.read_csv(io.StringIO(csv_text), nrows=MAX_EXCEL_ROWS)
            
            # Convert to structured format
            structured_data = df.to_dict('records')
            schema = df.columns.tolist()
            
            # Create text representation (limited for large files)
            if len(df) <= 1000:
                content_text = df.to_string()
            else:
                content_text = f"[Large CSV with {len(df)} rows - showing first 10 rows]\n"
                content_text += df.head(10).to_string()
            
            return {
                'type': 'structured',
                'content': content_text,
                'data': structured_data,
                'schema': schema,
                'row_count': len(df),
                'encoding': encoding,
                'extraction_method': 'pandas_csv_limited'
            }
            
        except Exception as e:
            logger.error(f"CSV extraction failed: {e}")
            raise
        finally:
            del content  # Help garbage collector
    
    def extract_text(self, content: bytes, metadata: Dict[str, Any]) -> Dict[str, Any]:
        """Extract plain text content."""
        try:
            # Detect encoding
            detected = chardet.detect(content[:10000])
            encoding = detected.get('encoding', 'utf-8')
            
            text = content.decode(encoding, errors='replace')
            
            return {
                'type': 'text',
                'content': text,
                'encoding': encoding,
                'extraction_method': 'text_decode'
            }
            
        except Exception as e:
            logger.error(f"Text extraction failed: {e}")
            raise
    
    def _detect_mime_type(self, filename: str, content: bytes) -> Optional[str]:
        """Detect MIME type from filename and content."""
        # Try from filename
        mime_type, _ = mimetypes.guess_type(filename)
        if mime_type:
            return mime_type
        
        # Try from content using python-magic if available
        try:
            import magic
            mime = magic.Magic(mime=True)
            return mime.from_buffer(content)
        except ImportError:
            logger.debug("python-magic not available for content-based MIME detection")
        except Exception as e:
            logger.debug(f"MIME detection failed: {e}")
        
        return None
    
    def _create_error_result(self, error_message: str) -> Dict[str, Any]:
        """Create an error result dictionary."""
        return {
            'type': 'error',
            'content': '',
            'error': error_message,
            'extraction_method': 'none'
        }
    
    def _log_memory_usage(self, context: str) -> None:
        """Log current memory usage."""
        try:
            import psutil
            process = psutil.Process()
            memory_info = process.memory_info()
            memory_mb = memory_info.rss / 1024 / 1024
            logger.info(f"Memory usage {context}: {memory_mb:.1f} MB")
        except ImportError:
            # psutil not available, use basic memory info
            import resource
            memory_kb = resource.getrusage(resource.RUSAGE_SELF).ru_maxrss
            # On Linux, ru_maxrss is in KB, on macOS it's in bytes
            if sys.platform == 'darwin':
                memory_mb = memory_kb / 1024 / 1024
            else:
                memory_mb = memory_kb / 1024
            logger.info(f"Memory usage {context}: {memory_mb:.1f} MB")
        except Exception as e:
            logger.debug(f"Could not log memory usage: {e}")